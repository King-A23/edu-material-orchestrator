from __future__ import annotations

import html
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer

from ..config import AppConfig
from ..utils.files import ensure_directory


EXPORT_FORMATS = {"docx", "pdf", "html"}
IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>.*?)]\((?P<path>.+?)\)$")


@dataclass(slots=True)
class HeadingBlock:
    level: int
    text: str


@dataclass(slots=True)
class ParagraphBlock:
    text: str


@dataclass(slots=True)
class ListBlock:
    ordered: bool
    items: list[str]


@dataclass(slots=True)
class QuoteBlock:
    lines: list[str]


@dataclass(slots=True)
class CodeBlock:
    language: str | None
    code: str


@dataclass(slots=True)
class ImageBlock:
    alt: str
    path: str


MarkdownBlock = HeadingBlock | ParagraphBlock | ListBlock | QuoteBlock | CodeBlock | ImageBlock


def parse_markdown(markdown_text: str) -> list[MarkdownBlock]:
    lines = markdown_text.splitlines()
    blocks: list[MarkdownBlock] = []
    index = 0

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip() or None
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(CodeBlock(language=language, code="\n".join(code_lines)))
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            blocks.append(HeadingBlock(level=max(1, min(level, 6)), text=text))
            index += 1
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            blocks.append(
                ImageBlock(
                    alt=image_match.group("alt").strip(),
                    path=image_match.group("path").strip(),
                )
            )
            index += 1
            continue

        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            blocks.append(QuoteBlock(lines=quote_lines))
            continue

        if _is_list_item(stripped):
            ordered = bool(re.match(r"^\d+\.\s+", stripped))
            items: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not candidate or not _is_list_item(candidate) or bool(re.match(r"^\d+\.\s+", candidate)) != ordered:
                    break
                items.append(_strip_list_marker(candidate))
                index += 1
            blocks.append(ListBlock(ordered=ordered, items=items))
            continue

        paragraph_lines: list[str] = []
        while index < len(lines):
            candidate_raw = lines[index]
            candidate = candidate_raw.strip()
            if not candidate:
                break
            if (
                candidate.startswith("```")
                or candidate.startswith("#")
                or candidate.startswith(">")
                or _is_list_item(candidate)
                or IMAGE_PATTERN.match(candidate)
            ):
                if paragraph_lines:
                    break
            paragraph_lines.append(candidate)
            index += 1
        blocks.append(ParagraphBlock(text=" ".join(paragraph_lines)))
    return blocks


def _is_list_item(line: str) -> bool:
    return bool(re.match(r"^[-*]\s+", line) or re.match(r"^\d+\.\s+", line))


def _strip_list_marker(line: str) -> str:
    line = re.sub(r"^[-*]\s+", "", line)
    line = re.sub(r"^\d+\.\s+", "", line)
    return line.strip()


def export_markdown_file(
    input_path: str | Path,
    targets: Iterable[str],
    output_path: str | Path | None = None,
    config: AppConfig | None = None,
) -> dict[str, str]:
    markdown_path = Path(input_path).resolve()
    markdown_text = markdown_path.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown_text)
    app_config = config or AppConfig.default()
    requested = []
    for item in targets:
        normalized = item.lower()
        if normalized not in EXPORT_FORMATS:
            raise ValueError(f"Unsupported export format: {item}")
        if normalized not in requested:
            requested.append(normalized)

    exported: dict[str, str] = {}
    for target in requested:
        if output_path is not None and len(requested) == 1:
            target_path = Path(output_path).resolve()
        else:
            target_path = markdown_path.with_suffix(f".{target}")
        ensure_directory(target_path.parent)
        if target == "html":
            render_markdown_to_html(markdown_path, blocks, target_path)
        elif target == "docx":
            render_markdown_to_docx(markdown_path, blocks, target_path)
        elif target == "pdf":
            render_markdown_to_pdf(markdown_path, blocks, target_path, page_size=app_config.render.pdf_page_size)
        exported[target] = str(target_path)
    return exported


def render_markdown_to_html(markdown_path: Path, blocks: list[MarkdownBlock], output_path: Path) -> Path:
    body_lines = ['<div class="document">']
    for block in blocks:
        if isinstance(block, HeadingBlock):
            body_lines.append(f"<h{block.level}>{_inline_html(block.text)}</h{block.level}>")
        elif isinstance(block, ParagraphBlock):
            body_lines.append(f"<p>{_inline_html(block.text)}</p>")
        elif isinstance(block, ListBlock):
            tag = "ol" if block.ordered else "ul"
            body_lines.append(f"<{tag}>")
            for item in block.items:
                body_lines.append(f"<li>{_inline_html(item)}</li>")
            body_lines.append(f"</{tag}>")
        elif isinstance(block, QuoteBlock):
            body_lines.append("<blockquote>")
            for line in block.lines:
                body_lines.append(f"<p>{_inline_html(line)}</p>")
            body_lines.append("</blockquote>")
        elif isinstance(block, CodeBlock):
            body_lines.append(f"<pre><code>{html.escape(block.code)}</code></pre>")
        elif isinstance(block, ImageBlock):
            image_path = _image_output_path(markdown_path, output_path.parent, block.path)
            body_lines.append(f'<figure><img src="{html.escape(image_path)}" alt="{html.escape(block.alt)}" /></figure>')
    body_lines.append("</div>")

    output_path.write_text(
        "\n".join(
            [
                "<!DOCTYPE html>",
                "<html lang=\"zh-CN\">",
                "<head>",
                "<meta charset=\"utf-8\" />",
                "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />",
                "<style>",
                "body { font-family: Georgia, 'Microsoft YaHei', serif; margin: 2rem auto; max-width: 860px; line-height: 1.7; color: #222; }",
                "h1,h2,h3,h4,h5,h6 { line-height: 1.3; }",
                "blockquote { border-left: 4px solid #999; margin: 1rem 0; padding: 0 1rem; color: #555; }",
                "pre { background: #f5f5f5; padding: 1rem; overflow-x: auto; }",
                "img { max-width: 100%; height: auto; }",
                "</style>",
                "</head>",
                "<body>",
                *body_lines,
                "</body>",
                "</html>",
            ]
        ),
        encoding="utf-8",
    )
    return output_path


def render_markdown_to_docx(markdown_path: Path, blocks: list[MarkdownBlock], output_path: Path) -> Path:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for block in blocks:
        if isinstance(block, HeadingBlock):
            document.add_heading(block.text, level=min(block.level, 4))
        elif isinstance(block, ParagraphBlock):
            document.add_paragraph(block.text)
        elif isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                document.add_paragraph(item, style=style)
        elif isinstance(block, QuoteBlock):
            for line in block.lines:
                paragraph = document.add_paragraph(line)
                paragraph.paragraph_format.left_indent = Inches(0.35)
                for run in paragraph.runs:
                    run.italic = True
        elif isinstance(block, CodeBlock):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.code)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif isinstance(block, ImageBlock):
            image_source = (markdown_path.parent / block.path).resolve()
            if image_source.exists():
                document.add_picture(str(image_source), width=Inches(5.5))
                if block.alt:
                    caption = document.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(block.alt)
                    caption_run.italic = True
            else:
                document.add_paragraph(f"[Missing image] {block.path}")
    document.save(output_path)
    return output_path


def render_markdown_to_pdf(
    markdown_path: Path,
    blocks: list[MarkdownBlock],
    output_path: Path,
    page_size: str = "A4",
) -> Path:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Quote",
            parent=styles["BodyText"],
            leftIndent=18,
            textColor=colors.HexColor("#555555"),
            italic=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            backColor=colors.HexColor("#F5F5F5"),
            leftIndent=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            alignment=1,
            italic=True,
            textColor=colors.HexColor("#555555"),
        )
    )
    story = []
    for block in blocks:
        if isinstance(block, HeadingBlock):
            style_name = "Heading1" if block.level == 1 else "Heading2" if block.level == 2 else "Heading3"
            story.append(Paragraph(_inline_html(block.text), styles[style_name]))
            story.append(Spacer(1, 8))
        elif isinstance(block, ParagraphBlock):
            story.append(Paragraph(_inline_html(block.text), styles["BodyText"]))
            story.append(Spacer(1, 8))
        elif isinstance(block, ListBlock):
            items = [ListItem(Paragraph(_inline_html(item), styles["BodyText"])) for item in block.items]
            story.append(
                ListFlowable(
                    items,
                    bulletType="1" if block.ordered else "bullet",
                    start="1",
                )
            )
            story.append(Spacer(1, 8))
        elif isinstance(block, QuoteBlock):
            for line in block.lines:
                story.append(Paragraph(_inline_html(line), styles["Quote"]))
            story.append(Spacer(1, 8))
        elif isinstance(block, CodeBlock):
            story.append(Preformatted(block.code, styles["CodeBlock"]))
            story.append(Spacer(1, 8))
        elif isinstance(block, ImageBlock):
            image_source = (markdown_path.parent / block.path).resolve()
            if image_source.exists():
                image_reader = ImageReader(str(image_source))
                width_px, height_px = image_reader.getSize()
                max_width = 360.0
                if width_px <= 0 or height_px <= 0:
                    pdf_image = PdfImage(str(image_source), width=max_width, hAlign="CENTER")
                else:
                    scale = min(1.0, max_width / float(width_px))
                    pdf_image = PdfImage(
                        str(image_source),
                        width=float(width_px) * scale,
                        height=float(height_px) * scale,
                        hAlign="CENTER",
                    )
                story.append(pdf_image)
                if block.alt:
                    story.append(Paragraph(_inline_html(block.alt), styles["Caption"]))
            else:
                story.append(Paragraph(_inline_html(f"[Missing image] {block.path}"), styles["BodyText"]))
            story.append(Spacer(1, 8))

    document = SimpleDocTemplate(str(output_path), pagesize=_pdf_page_size(page_size))
    document.build(story)
    return output_path


def _inline_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", escaped)
    return escaped


def _image_output_path(markdown_path: Path, output_parent: Path, raw_path: str) -> str:
    image_source = (markdown_path.parent / raw_path).resolve()
    return os.path.relpath(image_source, output_parent.resolve()).replace("\\", "/")


def _pdf_page_size(name: str):
    normalized = name.upper()
    if normalized == "LETTER":
        return LETTER
    return A4
